import os
import json
import sys
import time
import pathlib
import roman
import argparse
import logging
import subprocess
import traceback

from enum import Enum

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import RGBColor, Pt, Length
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.run import Run

# globals
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO, format='🐞 %(message)s')

DRY_RUN = False
UPDATE_LEGACY_BOXNOTES = False

# Track certain things globally while we recurse through the content
document = None
current_paragraph = None
current_path = None
use_table_cell_paragraph = False
current_table_cell = None

list_type = None
list_depths = {}
bullet_list_level = 0
ordered_list_level = 0
check_list_level = 0

in_bullet_list_item = False
in_ordered_list_item = False
in_check_list_item = False
is_check_list_item_checked = False

in_callout = False
callout_emoji = None
callout_bg_color = None

in_code_block = False

# Word limits highlights to a small pallete. Here's the mapping from Box options to Word options
highlight_map = {
    "#fdf0d1": WD_COLOR_INDEX.YELLOW,
    "#d4f3e6": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#ecd9fb": WD_COLOR_INDEX.PINK,
    "#fce6d1": WD_COLOR_INDEX.GRAY_50,
    "#ccdff7": WD_COLOR_INDEX.TURQUOISE,
    "#fbd7dd": WD_COLOR_INDEX.RED,
    "#e8e8e8": WD_COLOR_INDEX.GRAY_25,
}
heading_level_to_size_map = {"1": Pt(28), "2": Pt(20), "3": Pt(16)}

# Track which files failed to open or write for retry string output
failed_files = []

# Constants
INDENT = "    "
HOME = str(pathlib.Path.home())
OS_X_BOX_ROOT = os.path.sep.join([HOME, "Library", "CloudStorage", "Box-Box"])
WIN_BOX_ROOT = os.path.sep.join([HOME, "Box"])
BOX_ROOTS = [OS_X_BOX_ROOT, WIN_BOX_ROOT]  # Bit of a hack

# Enums
class Format(Enum):
    DOCX = 'docx'
    MD = 'md'
    HTML = 'html'

# Custom exceptions
class OldBoxNoteFormatError(Exception):
    pass

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('path', type=pathlib.Path, help="Path to a .boxnote or a directory")
    parser.add_argument('--format', default='docx', choices=[f.value for f in Format], help="Output format for converted boxnotes (default: docx)")
    parser.add_argument('--recursive', action='store_true', help="Convert all boxnotes within all subdirectories of the specified directory")
    parser.add_argument('--update_legacy_boxnotes', action='store_true', help="Update legacy Box Notes (created prior to August 2022) by opening them in your web browser (caution: may open a lot of browser tabs)")
    parser.add_argument('--dry-run', action='store_true', help="Perform a dry run (doesn't actually convert any files)")
    parser.add_argument('--debug', action='store_true', help='write debugging information to the console')
    args = parser.parse_args()

    global DRY_RUN
    if args.dry_run:
        print("🪄 performing a dry run, no files will be converted")
        DRY_RUN = True
    
    global UPDATE_LEGACY_BOXNOTES
    if args.update_legacy_boxnotes:
        print("⚙️ will update legacy boxnotes to the newer format")
        UPDATE_LEGACY_BOXNOTES = True
    
    if args.debug:
        logger.setLevel(logging.DEBUG)
        logger.debug("debug mode on")
    
    try:
        format = Format(value=args.format)
    except ValueError:
        print(f"invalid format specified: {args.format}")
        sys.exit(1)

    if format != Format.DOCX:
        print("⚠️ this script only supports docx, other formats will be added in the future")
        sys.exit(1)
    
    if args.path.is_file():
        convert_file(path=args.path, format=args.format)
    elif args.path.is_dir():
        convert_dir(path=args.path, recursively=args.recursive, format=format)


# Convert a directory of files
def convert_dir(path: pathlib.Path, recursively: bool, format: format) -> None:
    boxnotes = []
    failed_boxnotes = []
    if recursively is True:
        print(f"🔍 searching directory {path} and all subdirectories for boxnotes")
        boxnotes = [file_path for file_path in path.rglob(pattern='*.boxnote')]
    else:
        print(f"🔍 searching directory {path} for boxnotes")
        boxnotes = [file_path for file_path in path.glob(pattern='*.boxnote')]
    
    print(f'🔍 found {len(boxnotes)} boxnotes')

    for b in boxnotes:
        if not convert_file(path=b, format=format):
            failed_boxnotes.append(b)
    
    print()
    print("Conversion finished!")
    if len(failed_boxnotes) > 0:
        print("The following boxnotes failed to convert. If these were updgraded from the legacy Box Notes format (via --update_legacy_boxnotes), you may wish to re-run this script to convert them.")
        print('\n'.join([str(f) for f in failed_boxnotes]))
    else:
        print("All boxnotes were converted successfully.")

# Convert a single file
def convert_file(path: pathlib.Path, format: Format) -> None:
    if is_valid_path(path=path):
        if is_boxnote(path=path):
            return convert_boxnote(path=path, format=format)
        else:
            raise FileExistsError(f"file at path {path.absolute()} is not a boxnote")
    else:
        raise FileNotFoundError(f"file not found at path {path.absolute()}")

# Check that path exists and is a file
def is_valid_path(path: pathlib.Path) -> bool:
    return path.exists() and path.is_file()

# Check that path is a boxnote
def is_boxnote(path: pathlib.Path) -> bool:
    return path.suffix == '.boxnote'

def open_with_retry(path: pathlib.Path):
    timeouts = [1, 3, 5, 5, 10]
    num_tries = 0
    while True:
        try:
            with path.open() as f:
                bn_json = json.load(f)
            break
        except TimeoutError:
            if num_tries >= len(timeouts):
                print(f"⚠️ failed to open file {path}")
                raise TimeoutError(f"failed to open file {path} after multiple attempts")
            timeout = timeouts[num_tries]
            num_tries = num_tries + 1
            print(f"⏱️ waiting {timeout} seconds for file {path} to download...")
            time.sleep(timeout)
        except Exception as e:
            logger.debug(f"an unhandled exception occurred: {e}")
            print(traceback.format_exc())
            raise e
    
    return bn_json

# Open file, extract json and track failures
def convert_boxnote(path: pathlib.Path, format: Format) -> None:
    title = path.stem # remove ".boxnote"
    output_path = path.parent.joinpath(path.stem).with_suffix(suffix='.docx')

    if output_path.exists():
        print(f"⚠️ skipping {path.name}, converted file already exists → {output_path}")
    else:
        print(f"📄 converting {path.name} to {format} → {output_path}")

    if DRY_RUN:
        return
    
    global current_path
    current_path = path.absolute()

    try:
        bn_json = open_with_retry(path=path)
        logger.debug(f"loaded boxnote JSON:\n{repr(bn_json)}")
        return parse_boxnote_json(bn_json=bn_json, title=title, output_path=output_path.absolute())
    except OldBoxNoteFormatError:
        if UPDATE_LEGACY_BOXNOTES is True:
            print(f"⚙️ upgrading {path.name} → this boxnote is in an older fomat (prior to August 2022)")
            subprocess.call(args=["open", path.absolute()])
        else:
            print(f"⚠️ skipping {path.name} → this boxnote is in an older fomat (prior to August 2022)")
        return False
    except Exception as e:
        logger.debug(f"an unhandled exception occurred: {e}")
        print(traceback.format_exc())
        return False

# Parse JSON and build word document
def parse_boxnote_json(bn_json, title: str, output_path: pathlib.Path):
    global document, current_paragraph

    output = title + "\n"

    if "doc" in bn_json.keys():
        doc = bn_json["doc"]

        if "content" in doc.keys():
            document = Document()
            current_paragraph = document.add_paragraph()
            current_paragraph.add_run(text=title).font.size = Pt(points=32)
            content_objs = doc["content"]
            output = parse_contents(content_objs=content_objs, output=output)
            document.save(path_or_stream=output_path)
            
            if output_path.exists() and output_path.is_file() and output_path.stat().st_size > 0:
                return True
            else:
                print(f"⚠️ an error occured creating {output_path} 😢")
                return False
    else:
        raise OldBoxNoteFormatError

# Main entry into parsing is going to be the "content key"
def parse_contents(content_objs, output):
    for content_obj in content_objs:
        if "type" in content_obj.keys():
            output = parse_content_type(content_obj, content_obj["type"], output)
    return output


# Do different things depending on the obj_type we are parsing
def parse_content_type(content_obj, type, output):
    if type == "heading":
        return parse_heading_type(content_obj, output)
    elif type == "paragraph":
        return parse_paragraph_type(content_obj, output)
    elif type == "text":
        return parse_text_type(content_obj, output)
    elif type == "bullet_list":
        return parse_bullet_list_type(content_obj, output)
    elif type == "ordered_list":
        return parse_ordered_list_type(content_obj, output)
    elif type == "check_list":
        return parse_check_list_type(content_obj, output)
    elif type == "list_item":
        return parse_list_item_type(content_obj, output)
    elif type == "check_list_item":
        return parse_check_list_item_type(content_obj, output)
    elif type == "image":
        return parse_image_type(content_obj, output)
    elif type == "table":
        return parse_table_type(content_obj, output)
    elif type == "horizontal_rule":
        return parse_horizontal_rule_type(content_obj, output)
    elif type == "call_out_box":
        return parse_call_out_box(content_obj, output)
    elif type == "code_block":
        return parse_code_block(content_obj, output)
    elif type == "blockquote":
        return parse_blockquote(content_obj, output)
    return output


# Handle headings
def parse_heading_type(content_obj, output):
    global document, current_paragraph, heading_level_to_size_map
    current_paragraph = document.add_paragraph()
    heading_size = heading_level_to_size_map["1"]

    if "attrs" in content_obj.keys() and "level" in content_obj["attrs"]:
        heading_level = str(content_obj["attrs"]["level"])
        if heading_level in heading_level_to_size_map.keys():
            heading_size = heading_level_to_size_map[heading_level]
    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output)
        for run in current_paragraph.runs:
            run.font.size = heading_size
        return result + "\n"
    else:  # empty heading
        return output + "\n"


# Handle paragraphs
def parse_paragraph_type(content_obj, output):
    global document, current_paragraph, list_depths, ordered_list_level, is_check_list_item_checked, current_table_cell, use_table_cell_paragraph, in_callout, callout_emoji
    if (
        use_table_cell_paragraph
    ):  # happens due to table cells already having a paragraph
        current_paragraph = current_table_cell.paragraphs[0]
        use_table_cell_paragraph = False
    else:
        if current_table_cell is None:
            current_paragraph = document.add_paragraph()
        else:
            current_paragraph = current_table_cell.add_paragraph()

    if in_bullet_list_item:
        for _ in range(bullet_list_level - 1):
            current_paragraph.add_run(INDENT)
        current_paragraph.add_run("\u2022 ")

    elif in_ordered_list_item:
        for _ in range(ordered_list_level - 1):
            current_paragraph.add_run(INDENT)
        current_paragraph.add_run(
            get_ordered_list_char(ordered_list_level, list_depths[ordered_list_level])
            + " "
        )

    elif in_check_list_item:
        for _ in range(check_list_level - 1):
            current_paragraph.add_run(INDENT)
        current_paragraph.add_run(
            "\u2611 " if is_check_list_item_checked else "\u2610 "
        )
    elif in_callout:
        # Callout gets put in a table since that's close enough
        callout_table = document.add_table(1,1)
        cell = callout_table.cell(0,0)
        set_cell_background_color(cell, callout_bg_color)
        
        current_paragraph = cell.add_paragraph()
        current_paragraph.add_run(" " + callout_emoji + "  ")

    alignment = WD_ALIGN_PARAGRAPH.LEFT
    if "marks" in content_obj.keys():
        for mark_obj in content_obj["marks"]:
            if "type" in mark_obj:
                if (
                    mark_obj["type"] == "alignment"
                    and "attrs" in mark_obj.keys()
                    and "alignment" in mark_obj["attrs"].keys()
                ):
                    if mark_obj["attrs"]["alignment"] == "center":
                        alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif mark_obj["attrs"]["alignment"] == "right":
                        alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif mark_obj["attrs"]["alignment"] == "justify":
                        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    current_paragraph.alignment = alignment
    if "content" in content_obj.keys():
        return parse_contents(content_obj["content"], output) + "\n"
    else:  # empty paragraph
        return output + "\n"

# Handle text
def parse_text_type(content_obj, output):
    global current_paragraph, highlight_map, in_callout, callout_bg_color
    is_bold = False
    is_italic = False
    is_underline = False
    is_strikethrough = False
    font_color = None
    highlight_color = None
    font_size = None
    is_hyperlink = False
    url = None

    if "marks" in content_obj.keys():
        for mark_obj in content_obj["marks"]:
            if "type" in mark_obj:
                if mark_obj["type"] == "strong":
                    is_bold = True
                elif mark_obj["type"] == "em":
                    is_italic = True
                elif mark_obj["type"] == "underline":
                    is_underline = True
                elif mark_obj["type"] == "strikethrough":
                    is_strikethrough = True
                elif (
                    mark_obj["type"] == "font_size"
                    and "attrs" in mark_obj.keys()
                    and "size" in mark_obj["attrs"]
                ):
                    font_size = get_pt_from_em(mark_obj["attrs"]["size"])
                elif (
                    mark_obj["type"] == "font_color"
                    and "attrs" in mark_obj.keys()
                    and "color" in mark_obj["attrs"]
                ):
                    font_color = get_color_from_hex(mark_obj["attrs"]["color"])
                elif (
                    mark_obj["type"] == "highlight"
                    and "attrs" in mark_obj.keys()
                    and "color" in mark_obj["attrs"]
                ):
                    hex = mark_obj["attrs"]["color"]
                    if hex in highlight_map.keys():
                        highlight_color = highlight_map[hex]
                    else:
                        highlight_color = WD_COLOR_INDEX.YELLOW
                elif (
                    mark_obj["type"] == "link"
                    and "attrs" in mark_obj.keys()
                    and "href" in mark_obj["attrs"]
                ):
                    is_hyperlink = True
                    url = mark_obj["attrs"]["href"]

    if "text" in content_obj.keys():
        run = None
        if is_hyperlink:
            # Library doesn't support creating hyperlinks. Doing it manually
            part = current_paragraph.part
            r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)

            run = Run(OxmlElement("w:r"), current_paragraph)
            run.text = content_obj["text"]
            hyperlink.append(run._element)
            current_paragraph._p.append(hyperlink)
            font_color = RGBColor(26, 116, 186)
            is_underline = True
        else:
            run = current_paragraph.add_run(content_obj["text"])

        if is_bold:
            run.bold = True
        if is_italic:
            run.italic = True
        if is_underline:
            run.underline = True
        if is_strikethrough:
            run.font.strike = True
        if font_color is not None:
            run.font.color.rgb = font_color
        if highlight_color is not None:
            run.font.highlight_color = highlight_color
        if font_size is not None:
            run.font.size = font_size
        if in_code_block:
            run.font.name = "Courier"
        else:
          run.font.name = "Helvetica"

        if in_callout:
            set_run_background_color(run, callout_bg_color)

        return output + content_obj["text"]
    return output

# Set run background color
def set_run_background_color(run, bg_color):
    tag = run._r
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color)
    # run.font.size = Pt(14)
    tag.rPr.append(shd)

# Handle bullet lists
def parse_bullet_list_type(content_obj, output):
    global bullet_list_level, list_type
    list_type = "bullet"
    bullet_list_level += 1
    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output)
        bullet_list_level -= 1
        return result
    return output


# Handle ordered lists
def parse_ordered_list_type(content_obj, output):
    global ordered_list_level, list_depths, list_type
    list_type = "ordered"
    ordered_list_level += 1
    list_depths[ordered_list_level] = 0
    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output)
        ordered_list_level -= 1
        return result
    return output


# Handle checklists
def parse_check_list_type(content_obj, output):
    global check_list_level, list_type
    list_type = "check"
    check_list_level += 1
    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output)
        check_list_level -= 1
        return result
    return output


# Handle a checklist item
def parse_check_list_item_type(content_obj, output):
    global list_type, list_depths, check_list_level, in_check_list_item, is_check_list_item_checked
    if "attrs" in content_obj.keys() and "checked" in content_obj["attrs"].keys():
        if content_obj["attrs"]["checked"]:
            is_check_list_item_checked = True
        else:
            is_check_list_item_checked = False
    if "content" in content_obj.keys():
        padding = ""
        for _ in range(check_list_level):
            padding += "  "
        output += padding + ("\u2611 " if is_check_list_item_checked else "\u2610 ")
        in_check_list_item = True
        result = parse_contents(content_obj["content"], output)
        in_check_list_item = False

        if is_check_list_item_checked and len(current_paragraph.runs) > 1:
            is_check_found = False
            for run in current_paragraph.runs:
                if is_check_found:
                    run.font.strike = True
                elif run.text.find("\u2611") > -1:
                    is_check_found = True
        return result
    return output


# Handle either bullet or ordered list item
def parse_list_item_type(content_obj, output):
    global current_paragraph, in_bullet_list_item, in_ordered_list_item, list_type, list_depths, ordered_list_level
    if list_type == "ordered":
        list_depths[ordered_list_level] += 1
    if "content" in content_obj.keys():
        padding = ""
        if list_type == "bullet":
            for _ in range(bullet_list_level):
                padding += "  "
            output += padding + "- "
            in_bullet_list_item = True
        elif list_type == "ordered":
            for _ in range(ordered_list_level):
                padding += "  "
            output += (
                padding
                + get_ordered_list_char(
                    ordered_list_level, list_depths[ordered_list_level]
                )
                + " "
            )
            in_ordered_list_item = True
        result = parse_contents(content_obj["content"], output)
        in_bullet_list_item = False
        in_ordered_list_item = False
        return result
    return output


# Handle images
def parse_image_type(content_obj, output):
    global document, current_paragraph, current_path
    if "attrs" in content_obj.keys() and "fileName" in content_obj["attrs"].keys():
        image_file_name = content_obj["attrs"]["fileName"]
        image_file_path = get_image_path(image_file_name)
        if image_file_path is not None:
            if current_table_cell is None:
                document.add_picture(image_file_path, Length(in_to_emu(6)))
            else:
                current_paragraph.add_run().add_picture(
                    image_file_path, Length(in_to_emu(1))
                )
        else:
            current_paragraph.add_run("MISSING IMAGE: " + image_file_name)
        return output + "image: " + image_file_name


# Handle tables, unlike the other handlers, this one requires us
# to figure out the table structure before adding content
def parse_table_type(content_obj, output):
    global document, use_table_cell_paragraph, current_table_cell
    # Two passees through table data, once to determine dimensions and merged cells,
    # the other to fill the content
    r, c = get_table_dimensions(content_obj)

    # Tracking cell indexes in the table format is tricy due to the rowspans and colspans
    # Let's keep track of which cells have been allocated, so we can index more easily
    # A value of True indicates that this cell has already been "used" by a prior table cell's definition
    cell_tracking = get_cell_tracking_table(r, c)

    # Extract all table cells in order
    table_cell_objs = get_table_cell_objs(content_obj)

    # Get all the cells that need to be merged
    cell_merges = get_table_cell_merges(cell_tracking, table_cell_objs)

    # Add the table to the document
    table = document.add_table(r, c, "Table Grid")

    # Merge table cells that need merging
    merge_table_cells(table, cell_merges)

    # Table is finally ready, populate it.
    for table_cell_obj in table_cell_objs:
        parse_table_cell_type(table_cell_obj, table)

    current_table_cell = None
    use_table_cell_paragraph = False

    return output + "<Table " + str(r) + " x " + str(c) + ">"


# Handle a table cell
def parse_table_cell_type(table_cell_obj, table):
    global current_paragraph, current_table_cell, use_table_cell_paragraph

    row_idx = table_cell_obj["row_idx"]
    col_idx = table_cell_obj["col_idx"]

    current_table_cell = table.cell(row_idx, col_idx)
    # current_paragraph = table.cell(row_idx, col_idx).paragraphs[0]
    use_table_cell_paragraph = True

    parse_contents(table_cell_obj["content"], "")


# Merge table cells according to input
def merge_table_cells(table, cell_merges):
    merged_cells = []
    for r_idx, c_idx in cell_merges.keys():
        rowspan, colspan = cell_merges[(r_idx, c_idx)]
        cell = table.cell(r_idx, c_idx)
        other_cell = table.cell(r_idx + rowspan - 1, c_idx + colspan - 1)
        merged_cells.append(cell.merge(other_cell))
    return merged_cells


# Figure out which cells in the table are part of a merge
def get_table_cell_merges(cell_tracking, table_cell_objs):
    table_cell_objs_copy = table_cell_objs.copy()
    table_cell_objs_copy.reverse()  # reverse so we can pop our way throw the list
    # schema: (row index, column index) : (rowspan, colspan)
    cell_merges = {}
    # Iterate through the cell tracking table and write the index of each table cell
    for i_row in range(len(cell_tracking)):
        for i_col in range(len(cell_tracking[i_row])):
            if cell_tracking[i_row][i_col]:
                continue
            else:
                cell_tracking[i_row][i_col] = True
                table_cell_obj = table_cell_objs_copy.pop()
                table_cell_obj["row_idx"] = i_row
                table_cell_obj["col_idx"] = i_col
                rowspan = 1
                colspan = 1
                if "attrs" in table_cell_obj.keys():
                    if "rowspan" in table_cell_obj["attrs"].keys():
                        rowspan = table_cell_obj["attrs"]["rowspan"]
                    if "colspan" in table_cell_obj["attrs"].keys():
                        colspan = table_cell_obj["attrs"]["colspan"]

                if rowspan > 1 or colspan > 1:
                    cell_merges[(i_row, i_col)] = (rowspan, colspan)
                    for i_r in range(i_row, i_row + rowspan):
                        for i_c in range(i_col, i_col + colspan):
                            cell_tracking[i_r][i_c] = True
    return cell_merges


# Extract all the table cell objects in the JSON for this table
def get_table_cell_objs(content_obj):
    table_cell_objs = []
    for table_content_obj in content_obj["content"]:
        if (
            "type" in table_content_obj.keys()
            and table_content_obj["type"] == "table_row"
        ):
            # Iterate through table cells
            for row_content_obj in table_content_obj["content"]:
                if (
                    "type" in row_content_obj.keys()
                    and row_content_obj["type"] == "table_cell"
                ):
                    table_cell_objs.append(row_content_obj)
    return table_cell_objs


# Figure out the table's dimensions
def get_table_dimensions(content_obj):
    # total table_row types is the number of rows in table
    # number of columns is adding up the colspans of all cells in the first row
    # 1st level: single table
    # 2nd level: list of table_row
    # 3rd level: each table row has a list of table_cell
    row_objs = content_obj["content"]
    num_rows = len(row_objs)
    num_cols = 0

    # First row will give us our width by adding up the colspan
    for row_obj in row_objs[0]["content"]:
        if (
            "type" in row_obj.keys()
            and row_obj["type"] == "table_cell"
            and "attrs" in row_obj.keys()
        ):
            if "colspan" in row_obj["attrs"].keys():
                num_cols += row_obj["attrs"]["colspan"]

    return num_rows, num_cols


# Helper table used to determine the indexes of table cell objects
def get_cell_tracking_table(r, c):
    cell_tracking = []
    for _ in range(r):
        row = []
        for _ in range(c):
            row.append(False)
        cell_tracking.append(row)
    return cell_tracking

def parse_horizontal_rule_type(content_obj, output):
    global document, current_paragraph
    if current_paragraph is None:
      current_paragraph = document.add_paragraph()
    insert_horizontal_rule(current_paragraph)
    return output + "\n----------\n"

# Horizontal rule helper
def insert_horizontal_rule(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

# Parse call out box
def parse_call_out_box(content_obj, output):
    global callout_bg_color, in_callout, callout_emoji
    in_callout = True

    callout_emoji = ""
    if "attrs" in content_obj.keys():
      if "backgroundColor" in content_obj["attrs"].keys():
        callout_bg_color = (content_obj["attrs"]["backgroundColor"])[1:]
      if "emoji" in content_obj["attrs"].keys():
          callout_emoji = content_obj["attrs"]["emoji"]

    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output + callout_emoji + " ") + "\n"
        in_callout = False
        return result
    
    return output

# Parse call out box
def parse_code_block(content_obj, output):
    global document, current_paragraph, in_code_block
    in_code_block = True

    # Make a table for the code content, set it to a monospace font
    table = document.add_table(1, 1)
    cell = table.cell(0, 0)
    set_cell_background_color(cell, "ccdff7")
    current_paragraph = cell.paragraphs[0]

    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output) + "\n"
        in_code_block = False
        return result
    
    return output

# Parse call out box
def parse_blockquote(content_obj, output):
    if "content" in content_obj.keys():
        result = parse_contents(content_obj["content"], output) + "\n"
        in_code_block = False
        return result
    
    return output

# Get the appropriate list character(s) for this list item
# Box has numbers, then lowercase letters, then roman numerals
def get_ordered_list_char(ordered_list_level, list_depth):
    a = ordered_list_level % 3

    if a == 1:  # Numbers
        return str(list_depth) + "."

    elif a == 2:  # letters
        # convert to base-26, then convert each digit to the corresponding letter
        digits = []
        n = list_depth  # since we start at 1 instead of 0
        while n > 0:
            digits.insert(0, n % 26)
            n = n // 26

        char = ""
        for digit in digits:
            char += chr(digit + 96)
        return char + "."

    else:  # roman
        return roman.toRoman(list_depth).lower() + "."


# Convert from RGB Hex code to RGBColor object
def get_color_from_hex(hex):
    r = int(hex[1:3], 16)
    g = int(hex[3:5], 16)
    b = int(hex[5:7], 16)
    return RGBColor(r, g, b)



# Convert from em to pt
def get_pt_from_em(em):
    size = float(em[:-2]) // 0.083646
    return Pt(size)


# Search the current file's folder and its parents for the box image name specified
def get_image_path(img_file_name):
    global current_path

    # Try the current directory and all parents to see if the Box Note Image is available
    [base_path, tail] = os.path.split(current_path)
    file_name = tail[:-8]  # remove .boxnote
    extra_path = "Box Notes Images" + os.path.sep + file_name + " Images"

    while tail != "":
        img_file_path = (
            base_path + os.path.sep + extra_path + os.path.sep + img_file_name
        )

        if os.path.exists(img_file_path) and os.path.isfile(img_file_path):
            return img_file_path

        [base_path, tail] = os.path.split(base_path)
        if base_path in BOX_ROOTS:  # Actual break condition
            return None

    return None


# Convert inches to emu
def in_to_emu(inches):
    return inches * 914400

# Set a table cell's background color
def set_cell_background_color(cell, bgcolor):
  tc = cell._tc
  tc_props = tc.get_or_add_tcPr()
  tc_shading = OxmlElement('w:shd')
  tc_shading.set(qn('w:fill'), bgcolor)
  tc_props.append(tc_shading)


if __name__ == "__main__":
    main()
